import asyncio
import sys
from io import BytesIO

# Add src to path
sys.path.append("/app")

from docx import Document
from sqlalchemy import select

from src.infrastructure.uow import UnitOfWork
from src.modules.docs.models import DocsAIGenerationJob
from src.modules.docs.storage import DEFAULT_STORAGE_PROVIDER
from src.modules.files.models import File


async def check():
    job_id = "683d3fd8-9a2f-4112-92df-b2a43169129e"

    async with UnitOfWork() as uow:
        job = (
            await uow.session.execute(select(DocsAIGenerationJob).where(DocsAIGenerationJob.id == job_id))
        ).scalar_one_or_none()

        if not job:
            print("Job not found.")
            return

        file_obj = (await uow.session.execute(select(File).where(File.id == job.file_id))).scalar_one_or_none()

        if not file_obj:
            print("File not found.")
            return

        bucket = file_obj.s3_bucket
        key = file_obj.s3_key
        print(f"Reading from Bucket: {bucket}, Key: {key}")

        try:
            data = DEFAULT_STORAGE_PROVIDER.get_object_bytes(bucket=bucket, key=key)
            print(f"File size: {len(data)} bytes")

            doc = Document(BytesIO(data))
            print(f"Total Paragraphs: {len(doc.paragraphs)}")

            print("-" * 20)
            print("Document Content:")
            for p in doc.paragraphs:
                if p.text.strip():
                    # Check alignment (headers are centered)
                    align = "CENTER" if hasattr(p, "alignment") and p.alignment else "LEFT"
                    print(f"[{align}] {p.text}")
            print("-" * 20)

        except Exception as e:
            print(f"Error reading doc: {e}")


if __name__ == "__main__":
    asyncio.run(check())
