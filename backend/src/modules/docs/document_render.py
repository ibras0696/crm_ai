"""Рендеринг сгенерированного текста в TXT/DOCX/PDF форматы."""

from __future__ import annotations

import html
from io import BytesIO
from textwrap import wrap
from zipfile import ZIP_DEFLATED, ZipFile

from reportlab.pdfgen import canvas

from src.modules.docs.domain import FileType


def render_document_bytes(*, file_type: FileType, text: str, title: str | None) -> tuple[bytes, str, str]:
    """Собрать bytes-документ для указанного типа.

    Returns:
        tuple[payload, mime, extension]
    """
    normalized_title = str(title or "").strip() or "Документ"
    if file_type == FileType.TXT:
        return str(text or "").encode("utf-8"), "text/plain", "txt"
    if file_type == FileType.DOCX:
        return _build_docx_bytes(text=text, title=normalized_title), (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ), "docx"
    if file_type == FileType.PDF:
        return _build_pdf_bytes(text=text, title=normalized_title), "application/pdf", "pdf"
    raise ValueError("unsupported_file_type")


def _build_docx_bytes(*, text: str, title: str) -> bytes:
    """Собрать качественный DOCX из текста используя python-docx."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    
    # Настройка стилей по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Заголовок (если есть)
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph() # Отступ

    # Основной текст
    for line in (text or "").splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        doc.add_paragraph(line)

    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _build_pdf_bytes(*, text: str, title: str) -> bytes:
    """Собрать PDF из plain-текста (многостранично)."""
    stream = BytesIO()
    pdf = canvas.Canvas(stream, pagesize=(595, 842))
    pdf.setTitle(title)

    x = 56
    y = 800
    line_height = 16
    content_lines = []
    for raw in (text or "").splitlines() or [""]:
        wrapped = wrap(raw, width=90) or [""]
        content_lines.extend(wrapped)

    pdf.setFont("Helvetica", 11)
    for line in content_lines:
        if y < 60:
            pdf.showPage()
            pdf.setFont("Helvetica", 11)
            y = 800
        pdf.drawString(x, y, line)
        y -= line_height

    pdf.save()
    return stream.getvalue()
