"""DOCX processing for HTML conversion and document generation."""

from __future__ import annotations

import io
from dataclasses import dataclass
from html.parser import HTMLParser
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.shared import Inches

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph


@dataclass
class DocxProcessorResult:
    """Result of DOCX processing."""

    docx_bytes: bytes
    paragraphs_count: int
    word_count: int


class HTMLToDocxParser(HTMLParser):
    """Parse HTML and convert to DOCX elements."""

    def __init__(self, document: Document) -> None:
        super().__init__()
        self.document = document
        self.current_paragraph: Paragraph | None = None
        self.current_run: Any = None
        self.bold_stack: list[bool] = []
        self.italic_stack: list[bool] = []
        self.underline_stack: list[bool] = []
        self.list_stack: list[str] = []
        self.in_table = False

    def handle_starttag(self, tag: str, _attrs: list[tuple[str, str | None]]) -> None:
        """Handle opening HTML tags."""
        tag_lower = tag.lower()

        if tag_lower in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.current_paragraph = self.document.add_heading(level=int(tag_lower[1]))
            self.current_run = None
        elif tag_lower == "p":
            self.current_paragraph = self.document.add_paragraph()
            self.current_run = None
        elif tag_lower == "br":
            if self.current_paragraph:
                self.current_paragraph.add_run("\n")
        elif tag_lower in ("strong", "b"):
            self.bold_stack.append(True)
        elif tag_lower in ("em", "i"):
            self.italic_stack.append(True)
        elif tag_lower == "u":
            self.underline_stack.append(True)
        elif tag_lower == "ul":
            self.list_stack.append("bullet")
        elif tag_lower == "ol":
            self.list_stack.append("number")
        elif tag_lower == "li":
            if self.list_stack:
                list_type = self.list_stack[-1]
                style = "List Bullet" if list_type == "bullet" else "List Number"
                self.current_paragraph = self.document.add_paragraph(style=style)
                self.current_run = None
        elif tag_lower == "table":
            self.in_table = True
        elif tag_lower == "blockquote":
            self.current_paragraph = self.document.add_paragraph()
            if self.current_paragraph:
                self.current_paragraph.paragraph_format.left_indent = Inches(0.5)
            self.current_run = None

    def handle_endtag(self, tag: str) -> None:
        """Handle closing HTML tags."""
        tag_lower = tag.lower()

        if tag_lower in ("h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "blockquote"):
            self.current_paragraph = None
            self.current_run = None
        elif tag_lower in ("strong", "b"):
            if self.bold_stack:
                self.bold_stack.pop()
        elif tag_lower in ("em", "i"):
            if self.italic_stack:
                self.italic_stack.pop()
        elif tag_lower == "u":
            if self.underline_stack:
                self.underline_stack.pop()
        elif tag_lower in ("ul", "ol"):
            if self.list_stack:
                self.list_stack.pop()
        elif tag_lower == "table":
            self.in_table = False

    def handle_data(self, data: str) -> None:
        """Handle text content."""
        if self.in_table:
            return

        text = data.strip()
        if not text:
            return

        if not self.current_paragraph:
            self.current_paragraph = self.document.add_paragraph()

        run = self.current_paragraph.add_run(text)

        if self.bold_stack and self.bold_stack[-1]:
            run.bold = True
        if self.italic_stack and self.italic_stack[-1]:
            run.italic = True
        if self.underline_stack and self.underline_stack[-1]:
            run.underline = True

        self.current_run = run


class DocxProcessor:
    """DOCX processor for HTML conversion and document manipulation."""

    def html_to_docx(self, html_content: str) -> DocxProcessorResult:
        """Convert HTML to DOCX."""
        document = Document()
        parser = HTMLToDocxParser(document)
        parser.feed(html_content)

        output = io.BytesIO()
        document.save(output)
        output.seek(0)

        word_count = sum(len(paragraph.text.split()) for paragraph in document.paragraphs)

        return DocxProcessorResult(
            docx_bytes=output.read(),
            paragraphs_count=len(document.paragraphs),
            word_count=word_count,
        )

    def docx_to_html(self, docx_bytes: bytes) -> str:
        """Convert DOCX to HTML (basic conversion)."""
        document = Document(io.BytesIO(docx_bytes))
        html_parts: list[str] = []

        for paragraph in document.paragraphs:
            if not paragraph.text.strip():
                continue

            # Detect heading
            if paragraph.style.name.startswith("Heading"):
                level = paragraph.style.name.replace("Heading ", "")
                html_parts.append(f"<h{level}>{self._escape_html(paragraph.text)}</h{level}>")
            elif paragraph.style.name == "List Bullet" or paragraph.style.name == "List Number":
                html_parts.append(f"<li>{self._format_runs(paragraph)}</li>")
            else:
                html_parts.append(f"<p>{self._format_runs(paragraph)}</p>")

        return "\n".join(html_parts)

    def _format_runs(self, paragraph: Paragraph) -> str:
        """Format paragraph runs with HTML tags."""
        result: list[str] = []

        for run in paragraph.runs:
            text = self._escape_html(run.text)

            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"

            result.append(text)

        return "".join(result)

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )

    def create_empty_docx(self, title: str | None = None) -> bytes:
        """Create empty DOCX with optional title."""
        document = Document()

        if title:
            document.add_heading(title, level=1)

        document.add_paragraph("")

        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return output.read()

    def extract_text(self, docx_bytes: bytes) -> str:
        """Extract plain text from DOCX."""
        document = Document(io.BytesIO(docx_bytes))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    def get_word_count(self, docx_bytes: bytes) -> int:
        """Get word count from DOCX."""
        document = Document(io.BytesIO(docx_bytes))
        return sum(len(paragraph.text.split()) for paragraph in document.paragraphs)

    def extract_metadata(self, docx_bytes: bytes) -> dict[str, Any]:
        """Extract DOCX metadata."""
        document = Document(io.BytesIO(docx_bytes))
        core_properties = document.core_properties

        return {
            "title": core_properties.title or "",
            "author": core_properties.author or "",
            "subject": core_properties.subject or "",
            "keywords": core_properties.keywords or "",
            "comments": core_properties.comments or "",
            "created": core_properties.created.isoformat() if core_properties.created else None,
            "modified": core_properties.modified.isoformat() if core_properties.modified else None,
            "paragraphs": len(document.paragraphs),
            "word_count": self.get_word_count(docx_bytes),
        }

    def update_metadata(
        self,
        docx_bytes: bytes,
        metadata: dict[str, str],
    ) -> bytes:
        """Update DOCX metadata."""
        document = Document(io.BytesIO(docx_bytes))
        core_properties = document.core_properties

        if "title" in metadata:
            core_properties.title = metadata["title"]
        if "author" in metadata:
            core_properties.author = metadata["author"]
        if "subject" in metadata:
            core_properties.subject = metadata["subject"]
        if "keywords" in metadata:
            core_properties.keywords = metadata["keywords"]
        if "comments" in metadata:
            core_properties.comments = metadata["comments"]

        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return output.read()
